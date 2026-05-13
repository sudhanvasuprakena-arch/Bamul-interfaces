#!/usr/bin/env python3
"""Generate PO_AP.sql technical documentation as a Word document."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "docs/PO_AP_SQL_Technical_Documentation.docx"
DARK_BLUE  = RGBColor(0x1F, 0x4E, 0x79)
MID_BLUE   = RGBColor(0x2E, 0x75, 0xB6)
LIGHT_BLUE = RGBColor(0xBD, 0xD7, 0xEE)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
DARK_GREY  = RGBColor(0x26, 0x26, 0x26)
GREEN      = RGBColor(0x37, 0x86, 0x44)
CODE_BG    = RGBColor(0xF2, 0xF2, 0xF2)


def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '{:02X}{:02X}{:02X}'.format(rgb.red, rgb.green, rgb.blue))
    tcPr.append(shd)


def hdr(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h


def para(doc, text, bold=False, italic=False, size=11, color=DARK_GREY):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def bullet(doc, items, size=11):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(item)
        r.font.size = Pt(size)


def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)


def table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    # header row
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        set_cell_bg(cell, DARK_BLUE)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(10)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # data rows
    alt = RGBColor(0xF2, 0xF7, 0xFD)
    for ri, row in enumerate(rows):
        tr = t.add_row()
        bg = alt if ri % 2 == 0 else WHITE
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.text = str(val)
            set_cell_bg(cell, bg)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


def build_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── TITLE ─────────────────────────────────────────────────────────────────
    t = doc.add_heading('PO_AP.sql — Technical Documentation', 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = doc.add_paragraph('XXCUST_PO_AP_INTERFACE_PKG  |  Oracle EBS R12.2  |  Version 2.0')
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.runs[0].italic = True
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 1. PURPOSE
    # ══════════════════════════════════════════════════════════════════════════
    hdr(doc, '1. Purpose and Overview', 1)
    para(doc,
        'PO_AP.sql defines the Oracle PL/SQL package XXCUST_PO_AP_INTERFACE_PKG. '
        'Its job is to bridge the gap between the legacy Oracle EBS Purchasing module '
        'and the new Oracle EBS Accounts Payable module during a system migration. '
        'Because the new instance does not have a PO module, AP invoices cannot be '
        'created through the normal PO-matched invoice flow. This package reads '
        'confirmed receipt and return transactions directly from the legacy instance '
        'over a database link and creates the corresponding AP documents in the new '
        'instance via the standard AP Open Interface tables.')

    doc.add_paragraph()
    para(doc, 'The package has two public procedures:', bold=True)
    bullet(doc, [
        'run_receipt_interface  (Process A) — converts PO receipts into Standard AP Invoices.',
        'run_rtv_interface      (Process B) — converts post-invoice Return-to-Vendor transactions into AP Credit Memos.',
    ])

    para(doc, 'Core business rules enforced by the package:', bold=True)
    bullet(doc, [
        'NO RECEIPT = NO AP INVOICE. Only confirmed RECEIVE transactions trigger invoice creation.',
        'NET QUANTITY drives invoice amounts. Any RTV that occurs before the invoice is created is subtracted from the receipt quantity.',
        'POST-INVOICE RTV = CREDIT MEMO. An RTV that arrives after the invoice was already created generates a separate Credit Memo.',
        'NO DUPLICATES. Every receipt and every RTV is checked against the log table before processing.',
    ])

    # ══════════════════════════════════════════════════════════════════════════
    # 2. ARCHITECTURE
    # ══════════════════════════════════════════════════════════════════════════
    hdr(doc, '2. Architecture and Database Objects', 1)

    hdr(doc, '2.1 Database Link', 2)
    para(doc,
        'All legacy data is accessed through a database link named LEGACY_INSTANCE. '
        'Every table reference that ends with @legacy_instance reads from the old Oracle EBS '
        'schema. The new instance tables are accessed locally without a link.')

    hdr(doc, '2.2 Source Tables (Legacy Instance — read via DB link)', 2)
    table(doc,
        ['Table', 'Purpose'],
        [
            ('RCV_TRANSACTIONS',        'All receiving transactions — RECEIVE and RETURN TO VENDOR'),
            ('RCV_SHIPMENT_HEADERS',    'Receipt header — receipt number, date'),
            ('RCV_SHIPMENT_LINES',      'Receipt line details'),
            ('PO_HEADERS_ALL',          'PO header — PO number, vendor, currency, org'),
            ('PO_LINES_ALL',            'PO lines — item, description, purchase basis'),
            ('PO_LINE_LOCATIONS_ALL',   'PO shipment schedules — quantity billed, amount received'),
            ('PO_DISTRIBUTIONS_ALL',    'PO distributions — legacy CCID (account code)'),
            ('AP_SUPPLIERS',            'Supplier master — vendor name'),
            ('AP_SUPPLIER_SITES_ALL',   'Supplier site — site code'),
            ('HR_ALL_ORGANIZATION_UNITS','Operating unit name'),
            ('GL_CODE_COMBINATIONS',    'Legacy chart of accounts — segment values for COA mapping'),
            ('JAI_TAX_LINES_ALL',       'India localisation tax lines — CGST, SGST, IGST amounts'),
            ('MTL_SYSTEM_ITEMS_B',      'Item master — item description (used in account fallback)'),
        ],
        col_widths=[2.8, 4.5]
    )

    hdr(doc, '2.3 Target Tables (New Instance — written locally)', 2)
    table(doc,
        ['Table', 'Purpose'],
        [
            ('AP_INVOICES_INTERFACE',       'Staging table for AP invoice headers (Standard and Credit)'),
            ('AP_INVOICE_LINES_INTERFACE',  'Staging table for AP invoice lines (ITEM and TAX lines)'),
            ('XXCUST_PO_AP_INTERFACE_LOG',  'Custom log/control table — one row per processed transaction'),
            ('GL_CODE_COMBINATIONS',        'New COA — looked up to resolve new CCID from derived segments'),
        ],
        col_widths=[2.8, 4.5]
    )

    hdr(doc, '2.4 Custom Mapping Tables (New Instance)', 2)
    table(doc,
        ['Table', 'Purpose'],
        [
            ('MAPPING_DIVISION_SEGMENT',    'Maps legacy organization_id to new Division COA segment'),
            ('MAPPING_PRODUCT_SEGMENT',     'Maps legacy inventory_item_id to new Product COA segment'),
            ('MAPPING_DEPARTMENT_SEGMENT',  'Maps legacy cost-centre flex value to new Department segment'),
            ('MAPPING_ACCOUNT_SEGMENT',     'Maps legacy account flex value to new Account segment'),
            ('MAPPING_PRODUCT_PRE',         'Maps product code prefix to Account segment (fallback)'),
            ('MAPPING_ITEM_KEYWORD_ACCOUNT','Maps item description keywords to Account segment (fallback)'),
        ],
        col_widths=[2.8, 4.5]
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 3. CHART OF ACCOUNTS MAPPING
    # ══════════════════════════════════════════════════════════════════════════
    hdr(doc, '3. Chart of Accounts (COA) Mapping', 1)
    para(doc,
        'The new Oracle EBS instance uses a different COA structure from the legacy instance. '
        'Before any invoice line is inserted, the package derives a new Code Combination ID (CCID) '
        'by translating the legacy CCID into the new segment structure. This is done by the private '
        'function derive_new_ccid.')

    hdr(doc, '3.1 New COA Structure', 2)
    table(doc,
        ['Segment', 'Name', 'Source / Derivation'],
        [
            ('Segment 1', 'Entity',     'Fixed value: 01'),
            ('Segment 2', 'Division',   'MAPPING_DIVISION_SEGMENT  keyed on inventory organization_id'),
            ('Segment 3', 'Account',    'MAPPING_ACCOUNT_SEGMENT keyed on legacy Segment 4 (Account). Falls back to item-based lookup if no mapping found.'),
            ('Segment 4', 'Department', 'MAPPING_DEPARTMENT_SEGMENT keyed on legacy Segment 2 (Cost Centre)'),
            ('Segment 5', 'Product',    'MAPPING_PRODUCT_SEGMENT keyed on inventory_item_id'),
            ('Segment 6', 'TxnType',    'Derived from JAI tax lines: 1=Local (CGST/SGST), 2=Interstate (IGST), 0=Default'),
            ('Segment 7', 'Future 1',   'Fixed value: 000'),
            ('Segment 8', 'Future 2',   'Fixed value: 000'),
        ],
        col_widths=[1.0, 1.4, 5.0]
    )

    hdr(doc, '3.2 Account Segment Fallback Logic', 2)
    para(doc,
        'If the legacy account has no direct mapping in MAPPING_ACCOUNT_SEGMENT, '
        'the function get_account_segment_fallback is called. It tries two steps before '
        'returning the default account 513001 (General Consumables):')
    bullet(doc, [
        'Step 1 — Product prefix: looks up the item in MAPPING_PRODUCT_SEGMENT, takes the first two characters of the product code, and looks that prefix up in MAPPING_PRODUCT_PRE to find an account.',
        'Step 2 — Keyword matching: fetches the item description from MTL_SYSTEM_ITEMS_B and matches it against patterns in MAPPING_ITEM_KEYWORD_ACCOUNT, ordered by priority.',
        'Default — returns 513001 if neither step finds a match.',
    ])

    hdr(doc, '3.3 CCID Creation', 2)
    para(doc,
        'Once all eight segments are derived, derive_new_ccid looks up GL_CODE_COMBINATIONS in the '
        'new instance. If the combination does not exist yet, it calls FND_FLEX_EXT.GET_CCID to '
        'dynamically create it. If that also fails, the function returns -1 and the calling code '
        'substitutes a dummy CCID of 00000 so the invoice is still staged (to be corrected manually).')

    # ══════════════════════════════════════════════════════════════════════════
    # 4. PROCESS A
    # ══════════════════════════════════════════════════════════════════════════
    hdr(doc, '4. Process A — Receipt-to-Invoice', 1)
    para(doc,
        'Process A is the main invoice creation engine. It is called by run_receipt_interface '
        'and runs through five stages: Extract, Validate, Transform, Load, and Log.')

    hdr(doc, '4.1 What it does — plain English', 2)
    para(doc,
        'Every night (or on demand), Process A connects to the legacy Oracle instance and asks: '
        '"Which PO lines have been received but not yet fully billed?" For each such line, it '
        'calculates the net quantity that should be invoiced — gross received quantity minus any '
        'goods that were returned before the invoice was created. It then builds a Standard AP '
        'Invoice in the new instance for that net amount, including the correct tax lines, and '
        'records everything in the log table.')

    hdr(doc, '4.2 Stage 1 — Extract (cursor c_receipts)', 2)
    para(doc,
        'The cursor c_receipts queries the legacy instance and returns one row per RECEIVE '
        'transaction that meets all of the following conditions:')
    bullet(doc, [
        'transaction_type = RECEIVE and source_document_code = PO (only PO-sourced receipts).',
        'quantity_billed < quantity on the receipt (line is not already fully billed in legacy AP).',
        'Net quantity > 0 after subtracting all RETURN TO VENDOR quantities on the same PO line location. This is the Scenario A RTV netting filter — if all goods were returned before invoicing, the line is excluded entirely at cursor level.',
        'Optional filters: operating unit, receipt date range, PO number.',
    ])
    para(doc,
        'For each row the cursor also computes net_qty_received (gross minus all RTVs) and '
        'total_rtv_qty (sum of all RTVs) so the invoice amount and description can reflect '
        'the net position.')

    hdr(doc, '4.3 Stage 2 — Validate', 2)
    para(doc, 'Before loading, each receipt row passes through validate_receipt which checks:')
    table(doc,
        ['Rule', 'Check', 'Failure Action'],
        [
            ('BR-01', 'Net quantity and net amount must be > 0',                          'SKIPPED — logged with status SKIPPED'),
            ('BR-02', 'Vendor must exist in new AP_SUPPLIERS',                            'REJECTED — logged with reason'),
            ('BR-03', 'Receipt transaction not already processed as an invoice (dedup)',  'REJECTED — logged with reason'),
        ],
        col_widths=[0.8, 4.5, 2.5]
    )
    para(doc,
        'COA mapping (BR-04) is handled before validation by derive_new_ccid. If the CCID '
        'cannot be resolved, a dummy value 00000 is used and the invoice is still staged.')

    hdr(doc, '4.4 Stage 3 — Transform', 2)
    para(doc, 'The net invoice amount is computed as:')
    bullet(doc, [
        'Goods lines: net_qty_received x po_unit_price.',
        'Services lines (TEMP LABOR / FIXED PRICE): amount_received from PO_LINE_LOCATIONS_ALL.',
    ])
    para(doc,
        'The invoice number is generated as:  RCPT-{receipt_number}-L{po_line_num_padded_3_digits}')
    para(doc,
        'If the total_rtv_qty is greater than zero, the invoice description includes the text '
        '"(Net of RTV: {qty})" so there is a clear audit trail of the netting.')

    hdr(doc, '4.5 Stage 4 — Load', 2)
    para(doc, 'Three inserts are performed for each valid receipt:')
    bullet(doc, [
        '1. AP_INVOICES_INTERFACE — one header row. invoice_type_lookup_code = STANDARD. '
        'invoice_amount is initially set to the net item amount; it is updated after tax lines are added. '
        'The PO number is stored in attribute5 (not in the po_number column, because the new instance has no PO module).',

        '2. AP_INVOICE_LINES_INTERFACE — one ITEM line. line_type_lookup_code = ITEM. '
        'amount = net_qty x unit_price. quantity_invoiced and unit_price are populated for goods; '
        'NULL for services.',

        '3. AP_INVOICE_LINES_INTERFACE — one TAX line per non-zero tax component. '
        'Tax amounts are read from JAI_TAX_LINES_ALL on the legacy instance. '
        'Each tax component (CGST, SGST, IGST, TCS, CESS) becomes a separate line numbered '
        'as (po_line_num x 100) + sequence. After all tax lines are inserted, the invoice header '
        'amount is updated with UPDATE AP_INVOICES_INTERFACE SET invoice_amount = invoice_amount + total_tax.',
    ])

    hdr(doc, '4.6 Stage 5 — Log', 2)
    para(doc,
        'Every receipt — whether processed, rejected, skipped, or errored — gets a row in '
        'XXCUST_PO_AP_INTERFACE_LOG with the run_id, transaction_class = INVOICE, all key '
        'identifiers, quantities, amounts, and the final interface_status.')

    hdr(doc, '4.7 Process A Parameters', 2)
    table(doc,
        ['Parameter', 'Type', 'Default', 'Description'],
        [
            ('p_errbuf',            'OUT VARCHAR2', '—',    'Error message returned to caller / concurrent manager'),
            ('p_retcode',           'OUT NUMBER',   '—',    '0=Success, 1=Warning (rejections), 2=Fatal error'),
            ('p_operating_unit',    'IN VARCHAR2',  'NULL', 'Filter by operating unit name. NULL = all OUs'),
            ('p_receipt_date_from', 'IN VARCHAR2',  'NULL', 'Receipt date lower bound (DD-MON-YYYY). NULL = no lower limit'),
            ('p_receipt_date_to',   'IN VARCHAR2',  'NULL', 'Receipt date upper bound (DD-MON-YYYY). NULL = no upper limit'),
            ('p_po_number',         'IN VARCHAR2',  'NULL', 'Filter by PO number. NULL = all POs'),
            ('p_debug_mode',        'IN VARCHAR2',  'N',    'Y = write detailed netting and tax messages to DBMS_OUTPUT'),
        ],
        col_widths=[1.8, 1.3, 0.8, 3.8]
    )

    hdr(doc, '4.8 Process A — Execution Example', 2)
    code_block(doc,
        "DECLARE\n"
        "    l_errbuf  VARCHAR2(2000);\n"
        "    l_retcode NUMBER;\n"
        "BEGIN\n"
        "    XXCUST_PO_AP_INTERFACE_PKG.run_receipt_interface(\n"
        "        p_errbuf            => l_errbuf,\n"
        "        p_retcode           => l_retcode,\n"
        "        p_operating_unit    => NULL,\n"
        "        p_receipt_date_from => '01-JAN-2025',\n"
        "        p_receipt_date_to   => '31-DEC-2025',\n"
        "        p_po_number         => NULL,\n"
        "        p_debug_mode        => 'Y'\n"
        "    );\n"
        "    DBMS_OUTPUT.PUT_LINE('Return Code: ' || l_retcode);\n"
        "    DBMS_OUTPUT.PUT_LINE('Message    : ' || l_errbuf);\n"
        "END;\n"
        "/"
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 5. PROCESS B
    # ══════════════════════════════════════════════════════════════════════════
    hdr(doc, '5. Process B — RTV-to-Credit-Memo', 1)
    para(doc,
        'Process B handles the case where goods are returned to the vendor AFTER an AP Invoice '
        'has already been created by Process A. It generates a CREDIT type invoice (Credit Memo) '
        'in the new AP instance to offset the original Standard invoice.')

    hdr(doc, '5.1 What it does — plain English', 2)
    para(doc,
        'After Process A runs, Process B asks: "Are there any Return-to-Vendor transactions in '
        'the legacy system for which we already created an AP Invoice?" If yes, it creates a '
        'Credit Memo in the new AP instance for the returned quantity and amount, including '
        'reversed tax lines. If no invoice was ever created for that PO line (because the RTV '
        'happened before Process A ran), the RTV is simply skipped — it was already handled by '
        'the net quantity logic in Process A.')

    hdr(doc, '5.2 The Three RTV Scenarios', 2)
    table(doc,
        ['Scenario', 'When it happens', 'How it is handled'],
        [
            ('Scenario A',
             'RTV occurs BEFORE Process A runs. Net qty after RTV is still > 0.',
             'Process A nets the RTV off the receipt quantity. Invoice is created for net qty. Process B skips this RTV (logs as SKIPPED — Scenario A).'),
            ('Scenario A (full return)',
             'RTV occurs BEFORE Process A runs. Net qty after RTV = 0.',
             'Process A cursor excludes this line entirely (net qty filter). No invoice created. Process B skips this RTV.'),
            ('Scenario B / C',
             'RTV occurs AFTER Process A has already created an AP Invoice.',
             'Process B creates a CREDIT type invoice (Credit Memo) for the returned qty x unit price. Tax lines are also reversed.'),
        ],
        col_widths=[1.3, 3.2, 3.2]
    )

    hdr(doc, '5.3 Stage 1 — Extract (cursor c_rtvs)', 2)
    para(doc,
        'The cursor c_rtvs queries all RETURN TO VENDOR transactions from the legacy instance '
        'where transaction_type = RETURN TO VENDOR, source_document_code = PO, '
        'cancel_flag = N on the PO, and quantity > 0. The same optional filters '
        '(operating unit, date range, PO number) apply as in Process A.')

    hdr(doc, '5.4 Stage 2 — Scenario Detection', 2)
    para(doc,
        'For each RTV row, the function get_original_invoice_num is called. It queries '
        'XXCUST_PO_AP_INTERFACE_LOG to find whether a PROCESSED INVOICE record exists for '
        'the same PO line location. If it returns NULL, the RTV is Scenario A and is logged '
        'as SKIPPED. If it returns an invoice number, the RTV is Scenario B/C and processing continues.')

    hdr(doc, '5.5 Stage 3 — Validate', 2)
    para(doc, 'validate_rtv checks:')
    table(doc,
        ['Rule', 'Check', 'Failure Action'],
        [
            ('RTV-BR-01', 'RTV quantity and amount must be > 0',                                  'REJECTED'),
            ('RTV-BR-02', 'Original AP Invoice must exist (Scenario A already filtered upstream)', 'REJECTED'),
            ('RTV-BR-03', 'RTV not already processed as a Credit Memo (dedup)',                   'REJECTED'),
            ('RTV-BR-04', 'Credit amount must not exceed original invoice amount',                 'REJECTED — manual review required'),
        ],
        col_widths=[1.0, 4.5, 2.2]
    )

    hdr(doc, '5.6 Stage 4 — Load', 2)
    para(doc, 'Three inserts are performed for each valid RTV:')
    bullet(doc, [
        '1. AP_INVOICES_INTERFACE — one header row. invoice_type_lookup_code = CREDIT. '
        'invoice_amount = NEGATIVE (e.g. -74.66). The description includes the original invoice number for traceability.',

        '2. AP_INVOICE_LINES_INTERFACE — one ITEM line. amount = NEGATIVE credit amount. '
        'quantity_invoiced = RTV quantity (positive). Oracle AP handles the sign for CREDIT type.',

        '3. AP_INVOICE_LINES_INTERFACE — one TAX line per non-zero tax component from '
        'JAI_TAX_LINES on the legacy RTV transaction. Tax amounts are positive; Oracle AP '
        'reverses them automatically for CREDIT type invoices.',
    ])
    para(doc,
        'Credit Memo number format:  CM-{original_invoice_num}  '
        '(e.g. CM-RCPT-6797774-L075)')

    hdr(doc, '5.7 Process B Parameters', 2)
    table(doc,
        ['Parameter', 'Type', 'Default', 'Description'],
        [
            ('p_errbuf',         'OUT VARCHAR2', '—',    'Error message'),
            ('p_retcode',        'OUT NUMBER',   '—',    '0=Success, 1=Warning, 2=Fatal'),
            ('p_operating_unit', 'IN VARCHAR2',  'NULL', 'Filter by operating unit. NULL = all'),
            ('p_rtv_date_from',  'IN VARCHAR2',  'NULL', 'RTV date lower bound (DD-MON-YYYY)'),
            ('p_rtv_date_to',    'IN VARCHAR2',  'NULL', 'RTV date upper bound (DD-MON-YYYY)'),
            ('p_po_number',      'IN VARCHAR2',  'NULL', 'Filter by PO number. NULL = all'),
            ('p_debug_mode',     'IN VARCHAR2',  'N',    'Y = verbose logging'),
        ],
        col_widths=[1.8, 1.3, 0.8, 3.8]
    )

    hdr(doc, '5.8 Process B — Execution Example', 2)
    code_block(doc,
        "DECLARE\n"
        "    l_errbuf  VARCHAR2(2000);\n"
        "    l_retcode NUMBER;\n"
        "BEGIN\n"
        "    XXCUST_PO_AP_INTERFACE_PKG.run_rtv_interface(\n"
        "        p_errbuf         => l_errbuf,\n"
        "        p_retcode        => l_retcode,\n"
        "        p_operating_unit => NULL,\n"
        "        p_rtv_date_from  => '01-JAN-2025',\n"
        "        p_rtv_date_to    => '31-DEC-2025',\n"
        "        p_po_number      => NULL,\n"
        "        p_debug_mode     => 'Y'\n"
        "    );\n"
        "    DBMS_OUTPUT.PUT_LINE('Return Code: ' || l_retcode);\n"
        "    DBMS_OUTPUT.PUT_LINE('Message    : ' || l_errbuf);\n"
        "END;\n"
        "/"
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 6. PRIVATE FUNCTIONS REFERENCE
    # ══════════════════════════════════════════════════════════════════════════
    hdr(doc, '6. Private Functions Reference', 1)
    table(doc,
        ['Function', 'Returns', 'Purpose'],
        [
            ('log_message(p_message)',
             'void',
             'Writes a timestamped line to DBMS_OUTPUT. In production, replace with FND_FILE.PUT_LINE.'),
            ('get_net_received_qty(p_po_line_location_id)',
             'NUMBER',
             'Returns SUM(RECEIVE qty) - SUM(RTV qty) for a PO line location. Core of Scenario A netting.'),
            ('get_rtv_qty_for_line(p_po_line_location_id)',
             'NUMBER',
             'Returns total RTV quantity for a PO line location. Used by Process B for amount calculation.'),
            ('get_original_invoice_num(p_po_line_location_id, p_po_header_id)',
             'VARCHAR2',
             'Looks up XXCUST_PO_AP_INTERFACE_LOG for a PROCESSED INVOICE on this line. Returns invoice_num or NULL. Drives Scenario A vs B/C decision.'),
            ('validate_receipt(...)',
             'BOOLEAN',
             'Applies BR-01 to BR-03 business rules for Process A. Returns TRUE if valid.'),
            ('validate_rtv(...)',
             'BOOLEAN',
             'Applies RTV-BR-01 to RTV-BR-04 business rules for Process B. Returns TRUE if valid.'),
            ('get_division_segment(org_id)',
             'VARCHAR2',
             'Maps inventory organization_id to new Division COA segment via MAPPING_DIVISION_SEGMENT. Default: 01.'),
            ('get_product_segment(item_id)',
             'VARCHAR2',
             'Maps inventory_item_id to new Product COA segment via MAPPING_PRODUCT_SEGMENT. Default: 00000000.'),
            ('get_txn_type_segment(p_rcv_transaction_id)',
             'VARCHAR2',
             'Reads JAI_TAX_LINES_ALL to determine tax type. Returns 1 (CGST/SGST=Local), 2 (IGST=Interstate), 0 (default).'),
            ('get_department_segment(old_flex_value)',
             'VARCHAR2',
             'Maps legacy cost-centre flex value to new Department segment. Default: 00000.'),
            ('get_account_segment(old_flex_value)',
             'VARCHAR2',
             'Maps legacy account flex value to new Account segment. Returns NULL if not found (triggers fallback).'),
            ('get_account_segment_fallback(p_inventory_item_id)',
             'VARCHAR2',
             'Two-step fallback: product prefix lookup, then item description keyword match. Default: 513001.'),
            ('derive_new_ccid(...)',
             'NUMBER',
             'Orchestrates all segment derivations and looks up or creates the new CCID in GL_CODE_COMBINATIONS. Returns -1 on failure.'),
        ],
        col_widths=[2.5, 1.0, 4.2]
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 7. LOG TABLE
    # ══════════════════════════════════════════════════════════════════════════
    hdr(doc, '7. Log Table — XXCUST_PO_AP_INTERFACE_LOG', 1)
    para(doc,
        'Every transaction processed by either procedure is recorded in this table. '
        'It serves as both an audit trail and a control table for duplicate prevention.')
    table(doc,
        ['Column', 'Type', 'Description'],
        [
            ('log_id',                   'NUMBER (identity)', 'Primary key'),
            ('run_id',                   'VARCHAR2(50)',      'Unique run identifier. Format: INV-YYYYMMDDHH24MISS (Process A) or RTV-YYYYMMDDHH24MISS (Process B)'),
            ('transaction_class',        'VARCHAR2(20)',      'INVOICE (Process A) or CREDIT_MEMO (Process B)'),
            ('legacy_rcv_transaction_id','NUMBER',            'RCV_TRANSACTIONS.transaction_id from legacy — the RECEIVE or RTV transaction'),
            ('legacy_orig_rcv_txn_id',   'NUMBER',            'For RTVs: the original RECEIVE transaction_id'),
            ('legacy_po_number',         'VARCHAR2(20)',      'PO number from legacy'),
            ('legacy_po_line_num',       'NUMBER',            'PO line number'),
            ('receipt_quantity',         'NUMBER',            'Gross received quantity (RECEIVE transactions)'),
            ('rtv_quantity',             'NUMBER',            'Returned quantity (RETURN TO VENDOR transactions)'),
            ('net_quantity',             'NUMBER',            'receipt_quantity - rtv_quantity'),
            ('receipt_amount',           'NUMBER',            'Net invoiceable amount (after tax)'),
            ('invoice_num',              'VARCHAR2(50)',      'Generated invoice or credit memo number'),
            ('invoice_amount',           'NUMBER',            'Final invoice amount loaded to interface'),
            ('interface_status',         'VARCHAR2(20)',      'PROCESSED, REJECTED, SKIPPED, or ERROR'),
            ('rejection_reason',         'VARCHAR2(2000)',    'Populated for REJECTED, SKIPPED, and ERROR records'),
        ],
        col_widths=[2.2, 1.6, 3.9]
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 8. END-TO-END FLOW
    # ══════════════════════════════════════════════════════════════════════════
    hdr(doc, '8. End-to-End Execution Flow', 1)
    para(doc, 'The recommended daily execution sequence is:')
    table(doc,
        ['Step', 'Action', 'Who / How'],
        [
            ('1', 'Run Process A (run_receipt_interface)',
             'Node.js scheduler calls the PL/SQL procedure via oracledb. Captures DBMS_OUTPUT to log file.'),
            ('2', 'Run AP Open Interface Import for Process A',
             'Oracle EBS concurrent program: Payables Manager > Other > Import > Invoices. Source = XXCUST_PO_RECEIPT.'),
            ('3', 'Run Process B (run_rtv_interface)',
             'Node.js scheduler runs after Process A completes. Process B needs the log records from Process A to detect Scenario B/C.'),
            ('4', 'Run AP Open Interface Import for Process B',
             'Same concurrent program, same source. Imports the Credit Memos staged by Process B.'),
            ('5', 'Run reconciliation queries',
             'Query 4A (invoice recon), Query 4B (credit memo recon), Query 4C (net balance check).'),
            ('6', 'Log purge (optional)',
             'purge_log(p_days_to_keep => 90) removes PROCESSED records older than 90 days.'),
        ],
        col_widths=[0.5, 2.8, 4.4]
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 9. KNOWN ISSUES / NOTES
    # ══════════════════════════════════════════════════════════════════════════
    hdr(doc, '9. Known Issues and Notes', 1)
    bullet(doc, [
        'PO module not present in new instance: po_number and po_header_id columns are NOT populated in AP_INVOICES_INTERFACE and AP_INVOICE_LINES_INTERFACE. The PO number is stored in attribute5 instead.',
        'CCID fallback: If derive_new_ccid cannot resolve a valid CCID, it substitutes 00000. These invoices will fail AP validation and must be corrected manually before the AP Import can process them.',
        'Tax table reference: The tax query uses ja.jai_tax_lines@legacy_instance in Process B (schema prefix ja) vs apps.jai_tax_lines_all@legacy_instance in Process A. Ensure the correct schema and synonym are accessible via the DB link.',
        'Services lines: For purchase_basis = TEMP LABOR or FIXED PRICE, quantity_invoiced and unit_price are set to NULL and amount_received is used instead of qty x price.',
        'Pending fields: TERMS_ID, TERMS_NAME, TERM_DATE, and PAYMENT_METHOD_LOOKUP_CODE are not yet populated in the interface tables. These need to be added for complete payment terms handling.',
        'transaction_type filter: The cursor comment notes it should be DELIVER not RECEIVE for some receipt flows. Verify with the legacy receiving setup.',
    ])

    doc.save(OUT)
    print(f'Saved: {OUT}')


if __name__ == '__main__':
    build_doc()
